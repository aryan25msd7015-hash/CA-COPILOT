"""Deterministic engines powering the advanced automation modules."""
import io
import re
from calendar import monthrange
from collections import defaultdict
from datetime import date, datetime, timedelta
from html import escape
from statistics import mean

import openpyxl
from docx import Document as DocxDocument
from docx.shared import Inches
from openpyxl.styles import Font, PatternFill
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from reportlab.lib import colors


FILING_MATRIX = {
    "proprietorship": ["GSTR1", "GSTR3B", "ITR", "TDS_26Q"],
    "partnership": ["GSTR1", "GSTR3B", "ITR", "TDS_26Q", "FORM_3CB"],
    "pvt_ltd": ["GSTR1", "GSTR3B", "TDS_24Q", "TDS_26Q", "ROC_AOC4", "ROC_MGT7", "ADV_TAX", "FORM_3CA"],
    "llp": ["GSTR1", "GSTR3B", "TDS_26Q", "ROC_LLP8", "ROC_LLP11", "ADV_TAX"],
    "trust": ["ITR_7", "FCRA_IF_APPLICABLE"],
}

FILING_NAMES = {
    "GSTR1": "GSTR-1", "GSTR3B": "GSTR-3B", "ITR": "Income Tax Return",
    "ITR_7": "ITR-7", "TDS_24Q": "TDS 24Q", "TDS_26Q": "TDS 26Q",
    "FORM_3CA": "Tax Audit Form 3CA", "FORM_3CB": "Tax Audit Form 3CB",
    "ROC_AOC4": "ROC AOC-4", "ROC_MGT7": "ROC MGT-7", "ROC_LLP8": "LLP Form 8",
    "ROC_LLP11": "LLP Form 11", "ADV_TAX": "Advance Tax", "FCRA_IF_APPLICABLE": "FCRA Return",
}

CERTIFICATE_TYPES = {
    "net_worth": ("Certificate of Net Worth", ["net_worth_fy1", "net_worth_fy2", "net_worth_fy3", "tangible_net_worth", "total_liabilities"]),
    "turnover": ("Certificate of Turnover", ["turnover_fy1", "turnover_fy2", "turnover_fy3", "export_turnover"]),
    "working_capital": ("Working Capital Certificate", ["current_assets", "current_liabilities", "net_working_capital", "current_ratio"]),
    "end_use_of_funds": ("End Use of Funds Certificate", ["loan_amount", "utilized_amount", "balance_unspent", "utilization_details"]),
    "solvency": ("Solvency Certificate", ["total_assets", "total_liabilities", "net_worth", "solvency_ratio"]),
    "fund_utilization": ("Fund Utilization Certificate", ["funds_received", "funds_utilized", "unutilized_balance", "purpose"]),
}

ACTIVITY_BILLING_MAP = {
    "reconciliation_run": {"type": "GST Work", "billable": True},
    "notice_draft": {"type": "Tax Advisory", "billable": True},
    "certificate_gen": {"type": "Certificate", "billable": True},
    "document_review": {"type": "Audit Work", "billable": True},
    "export": {"type": "Report Prep", "billable": True},
    "upload": {"type": "Admin", "billable": False},
    "query_run": {"type": "Analysis", "billable": True},
    "secretarial_gen": {"type": "Secretarial", "billable": True},
    "lease_schedule": {"type": "Audit Work", "billable": True},
    "rfp_bid": {"type": "Business Development", "billable": False},
}


def month_period(today=None):
    return (today or date.today()).strftime("%Y-%m")


def compute_deadline(filing_type, period):
    year, month = map(int, period.split("-"))
    next_month = date(year + (month == 12), 1 if month == 12 else month + 1, 1)
    offsets = {
        "GSTR1": 11, "GSTR3B": 20, "TDS_24Q": 31, "TDS_26Q": 31,
        "ADV_TAX": 15, "ITR": 31, "ITR_7": 31, "FORM_3CA": 30, "FORM_3CB": 30,
        "ROC_AOC4": 30, "ROC_MGT7": 60, "ROC_LLP8": 30, "ROC_LLP11": 30,
        "FCRA_IF_APPLICABLE": 31,
    }
    if filing_type in ("GSTR1", "GSTR3B"):
        return next_month + timedelta(days=offsets[filing_type] - 1)
    if filing_type.startswith("TDS_"):
        return date(year, month, monthrange(year, month)[1]) + timedelta(days=31)
    if filing_type == "ADV_TAX":
        return date(year, month, 15)
    return date(year, month, monthrange(year, month)[1]) + timedelta(days=offsets.get(filing_type, 30))


def deadline_risk_score(deadline, data_received, late_count, has_open_notice, health_score, today=None):
    days_left = (deadline - (today or date.today())).days
    score = 3 if days_left <= 2 else 2 if days_left <= 5 else 1 if days_left <= 10 else 0
    score += 0 if data_received else 2.5
    score += 2 if late_count >= 3 else 1 if late_count >= 1 else 0
    score += 1.5 if has_open_notice else 0
    score += 1 if health_score < 50 else 0
    return round(min(score, 10), 1)


def parse_udyam_certificate(text):
    udyam = re.search(r"UDYAM-[A-Z]{2}-\d{2}-\d{7}", text or "", re.I)
    category = re.search(r"\b(MICRO|SMALL|MEDIUM)\s+ENTERPRISE\b", text or "", re.I)
    gstin = re.search(r"\b[0-9]{2}[A-Z]{5}[0-9]{4}[A-Z][1-9A-Z]Z[0-9A-Z]\b", text or "")
    name = re.search(r"Name of Enterprise\s*[:\-]\s*([^\r\n]+)", text or "", re.I)
    if not udyam or not category:
        raise ValueError("Udyam number and enterprise category are required")
    return {
        "udyam_reg_no": udyam.group(0).upper(),
        "udyam_category": category.group(1).lower(),
        "vendor_gstin": gstin.group(0) if gstin else None,
        "vendor_name": name.group(1).strip() if name else "MSME Vendor",
    }


def get_fy(day):
    year = day.year if day.month >= 4 else day.year - 1
    return f"{year}-{str(year + 1)[-2:]}"


def msme_violation_values(invoice_date, amount, payment_date=None, today=None):
    due = invoice_date + timedelta(days=45)
    reference = payment_date or today or date.today()
    days_overdue = max((reference - due).days, 0)
    violated = payment_date is None and days_overdue > 0 or payment_date is not None and payment_date > due
    return {
        "violated": violated,
        "due_date": due,
        "days_overdue": days_overdue if violated else 0,
        "disallowance_amount": round(float(amount), 2) if violated else 0,
        "interest_amount": round(float(amount) * 0.045 * (days_overdue / 365), 2) if violated else 0,
    }


def analyze_stock(items, rules, today=None):
    cutoff = int(rules.get("stock_age_cutoff_days", 180))
    margin = float(rules.get("stock_margin", 0.25))
    current = today or date.today()
    rows, gross, eligible = [], 0.0, 0.0
    for item in items:
        value = float(item.stock_value or 0)
        days = (current - item.last_movement_date).days if item.last_movement_date else cutoff + 1
        is_eligible = days <= cutoff
        gross += value
        eligible += value if is_eligible else 0
        rows.append({"sku": item.sku, "value": value, "days_since_movement": days, "eligible": is_eligible})
    return {
        "gross_stock": round(gross, 2), "eligible_stock": round(eligible, 2),
        "stock_dp": round(eligible * (1 - margin), 2),
        "ineligible_value": round(gross - eligible, 2), "items": rows,
    }


def analyze_debtors(items, rules, today=None):
    cutoff = int(rules.get("debtor_age_cutoff_days", 90))
    margin = float(rules.get("debtor_margin", 0.25))
    current = today or date.today()
    rows, gross, eligible, at_risk = [], 0.0, 0.0, 0
    for item in items:
        value = float(item.outstanding or 0)
        age = (current - item.invoice_date).days
        history = float(item.payment_history_score or 0)
        risk = age >= max(cutoff - 15, 0) or history < 75
        is_eligible = age <= cutoff and history >= 50
        gross += value
        eligible += value if is_eligible else 0
        at_risk += int(risk)
        rows.append({"debtor": item.debtor_name, "outstanding": value, "age_days": age, "at_risk": risk, "eligible": is_eligible})
    return {
        "gross_debtors": round(gross, 2), "eligible_debtors": round(eligible, 2),
        "debtor_dp": round(eligible * (1 - margin), 2), "at_risk_count": at_risk,
        "ineligible_value": round(gross - eligible, 2), "items": rows,
    }


def compute_drawing_power(stock, debtors, creditors, rules, sanctioned_limit):
    raw = stock["stock_dp"] + debtors["debtor_dp"]
    if rules.get("creditor_deduction", True):
        raw -= float(creditors or 0)
    return round(max(0, min(raw, float(sanctioned_limit))), 2)


def extract_certificate_fields(text, cert_type):
    if cert_type not in CERTIFICATE_TYPES:
        raise ValueError("Unsupported certificate type")
    fields = {}
    for field in CERTIFICATE_TYPES[cert_type][1]:
        label = field.replace("_", r"[\s_]+")
        match = re.search(rf"{label}\s*[:\-]?\s*(?:INR|Rs\.?|₹)?\s*([\d,]+(?:\.\d+)?)", text or "", re.I)
        fields[field] = float(match.group(1).replace(",", "")) if match else None
    if cert_type == "working_capital":
        assets, liabilities = fields.get("current_assets"), fields.get("current_liabilities")
        if assets is not None and liabilities is not None:
            fields["net_working_capital"] = round(assets - liabilities, 2)
            fields["current_ratio"] = round(assets / liabilities, 2) if liabilities else None
    if cert_type == "solvency":
        assets, liabilities = fields.get("total_assets"), fields.get("total_liabilities")
        if assets is not None and liabilities is not None:
            fields["net_worth"] = round(assets - liabilities, 2)
            fields["solvency_ratio"] = round(assets / liabilities, 2) if liabilities else None
    return fields


def validate_certificate_fields(fields, reference=None):
    issues = []
    reference = reference or {}
    for key, official in reference.items():
        extracted = fields.get(key)
        if extracted is None or not official:
            continue
        variance = abs(float(extracted) - float(official)) / abs(float(official)) * 100
        if variance > 2:
            issues.append({"field": key, "extracted": extracted, "official_value": official, "variance_pct": round(variance, 1)})
    missing = [key for key, value in fields.items() if value in (None, "")]
    return {"valid": not issues and not missing, "issues": issues, "missing_fields": missing}


def build_certificate_docx(title, client, fields, validation):
    doc = DocxDocument()
    doc.add_heading(title, 0)
    doc.add_paragraph(f"To whomsoever it may concern", style="Subtitle")
    doc.add_paragraph(
        f"We certify that the following particulars relating to {client.name} "
        f"(GSTIN: {client.gstin or 'Not available'}) have been verified from the records produced before us."
    )
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Particular"
    table.rows[0].cells[1].text = "Value"
    for key, value in fields.items():
        cells = table.add_row().cells
        cells[0].text = key.replace("_", " ").title()
        cells[1].text = f"INR {value:,.2f}" if isinstance(value, (int, float)) else str(value or "")
    doc.add_paragraph(f"Validation status: {'Verified' if validation.get('valid') else 'Review required'}")
    doc.add_paragraph(f"Date: {date.today().strftime('%d %B %Y')}")
    doc.add_paragraph("For the Chartered Accountant Firm\nAuthorized Signatory\nMembership No.: __________")
    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


def _resolution_items(transcript):
    chunks = [chunk.strip(" .") for chunk in re.split(r"[\r\n]+|(?<=[.!?])\s+", transcript or "") if len(chunk.strip()) > 10]
    return chunks[:8] or ["The matters placed before the meeting were considered and approved."]


def generate_secretarial_document(doc_type, client, transcript, data):
    meeting_date = data.get("meeting_date") or str(date.today())
    meeting_type = data.get("meeting_type", "Board Meeting")
    chairman = data.get("chairman", "Chairperson")
    directors = data.get("directors_present", ["Director 1", "Director 2"])
    items = _resolution_items(transcript)
    resolutions = [
        {"number": index + 1, "title": item[:80], "text": f"RESOLVED THAT {item.rstrip('.')} and the same was passed with requisite majority.", "passed_by": "requisite majority"}
        for index, item in enumerate(items)
    ]
    structured = {
        "company_name": client.name, "cin": client.cin or "", "meeting_date": meeting_date,
        "meeting_type": meeting_type, "venue": data.get("venue", client.registered_office or "Registered Office"),
        "chairman": chairman, "directors_present": directors, "directors_absent": data.get("directors_absent", []),
        "agenda_items": items, "resolutions": resolutions, "special_notes": data.get("special_notes"),
    }
    if doc_type == "agm_notice":
        generated = f"NOTICE OF ANNUAL GENERAL MEETING\n\nNotice is hereby given that the AGM of {client.name} will be held on {meeting_date} at {structured['venue']}.\n\nAGENDA\n" + "\n".join(f"{i + 1}. {item}" for i, item in enumerate(items))
    elif doc_type in ("mgt7", "aoc4"):
        structured.update(data)
        generated = f"{doc_type.upper()} DATA SHEET\n\n" + "\n".join(f"{key.replace('_', ' ').title()}: {value}" for key, value in structured.items() if key not in ("resolutions", "agenda_items"))
    else:
        generated = (
            f"MINUTES OF {meeting_type.upper()}\n{client.name}\nCIN: {client.cin or 'Not available'}\n"
            f"Date: {meeting_date}\nVenue: {structured['venue']}\nChairman: {chairman}\n\n"
            + "\n\n".join(f"{r['number']}. {r['title']}\n{r['text']}" for r in resolutions)
        )
    xml = None
    if doc_type in ("board_minutes", "mgt14"):
        xml = (
            '<?xml version="1.0" encoding="UTF-8"?>\n<MGT14>'
            f"<CompanyDetails><CIN>{escape(client.cin or '')}</CIN><CompanyName>{escape(client.name)}</CompanyName></CompanyDetails>"
            f"<FilingDetails><MeetingDate>{escape(str(meeting_date))}</MeetingDate><MeetingType>{escape(meeting_type)}</MeetingType>"
            + "".join(
                f"<Resolution><SerialNo>{r['number']}</SerialNo><ResolutionType>Ordinary</ResolutionType>"
                f"<ResolutionText>{escape(r['text'])}</ResolutionText></Resolution>"
                for r in resolutions
            ) + "</FilingDetails></MGT14>"
        )
    return structured, generated, xml


def build_text_docx(title, body):
    doc = DocxDocument()
    doc.add_heading(title, 0)
    for block in body.split("\n\n"):
        doc.add_paragraph(block)
    stream = io.BytesIO()
    doc.save(stream)
    return stream.getvalue()


def extract_lease_data(text, overrides=None):
    overrides = overrides or {}
    def number(pattern, default=0):
        found = re.search(pattern, text or "", re.I)
        return float(found.group(1).replace(",", "")) if found else default
    def integer(pattern, default=0):
        return int(number(pattern, default))
    commencement = re.search(r"(?:commencement|start)\s+date\s*[:\-]?\s*(\d{4}-\d{2}-\d{2})", text or "", re.I)
    data = {
        "commencement_date": commencement.group(1) if commencement else str(date.today()),
        "lease_term_months": integer(r"lease\s+term\s*[:\-]?\s*(\d+)\s*months?", 36),
        "base_rent_monthly": number(r"(?:monthly\s+rent|base\s+rent)\s*[:\-]?\s*(?:INR|Rs\.?|₹)?\s*([\d,]+(?:\.\d+)?)", 100000),
        "rent_free_period_months": integer(r"rent[\s-]*free\s+period\s*[:\-]?\s*(\d+)\s*months?", 0),
        "security_deposit": number(r"security\s+deposit\s*[:\-]?\s*(?:INR|Rs\.?|₹)?\s*([\d,]+(?:\.\d+)?)", 0),
        "incremental_borrowing_rate_pct": number(r"(?:IBR|incremental\s+borrowing\s+rate)\s*[:\-]?\s*([\d.]+)\s*%", 9),
        "escalation_clauses": [], "renewal_options": [], "termination_option": {},
    }
    escalation = re.search(r"(?:escalat\w*|increase)\s*(?:by)?\s*([\d.]+)\s*%.*?(?:month|after)\s*(\d+)", text or "", re.I)
    if escalation:
        data["escalation_clauses"] = [{"effective_from_month": int(escalation.group(2)), "rate_pct": float(escalation.group(1)), "type": "fixed"}]
    data.update(overrides)
    payments = []
    for month in range(1, int(data["lease_term_months"]) + 1):
        payment = 0 if month <= int(data.get("rent_free_period_months") or 0) else float(data["base_rent_monthly"])
        for clause in data.get("escalation_clauses") or []:
            if month >= int(clause.get("effective_from_month", 10**9)) and clause.get("type") == "fixed":
                payment *= 1 + float(clause.get("rate_pct", 0)) / 100
        payments.append({"month": month, "payment": round(payment, 2)})
    data["lease_payments_schedule"] = payments
    return data


def compute_lease_schedule(lease):
    payments = [float(row["payment"]) for row in lease.get("lease_payments_schedule", [])]
    if not payments:
        raise ValueError("Lease payment schedule is empty")
    rate = float(lease.get("incremental_borrowing_rate_pct") or 0) / 100 / 12
    liability = sum(payment / ((1 + rate) ** (index + 1)) for index, payment in enumerate(payments))
    initial_liability = liability
    rou_asset = initial_liability
    depreciation = rou_asset / len(payments)
    result = []
    for index, payment in enumerate(payments, 1):
        interest = liability * rate
        principal = payment - interest
        liability -= principal
        rou_asset -= depreciation
        result.append({
            "month": index, "payment": round(payment, 2), "interest_expense": round(interest, 2),
            "principal": round(principal, 2), "lease_liability": round(max(liability, 0), 2),
            "rou_asset": round(max(rou_asset, 0), 2),
        })
    return {"initial_lease_liability": round(initial_liability, 2), "initial_rou_asset": round(initial_liability, 2), "schedule": result}


def check_rfp_eligibility(text, credentials):
    criteria = []
    years_match = re.search(r"(?:minimum|at least)\s+(\d+)\s+years", text or "", re.I)
    required_years = int(years_match.group(1)) if years_match else 0
    firm_years = max(date.today().year - int(credentials.founding_year or date.today().year), 0)
    if required_years:
        criteria.append({"criterion": "Firm experience", "required_value": f"{required_years} years", "firm_value": f"{firm_years} years", "eligible": firm_years >= required_years, "gap": None if firm_years >= required_years else f"Requires {required_years - firm_years} more years"})
    turnover_match = re.search(r"(?:turnover|fee receipts).*?(?:INR|Rs\.?|₹)?\s*([\d,]+)", text or "", re.I)
    if turnover_match:
        required = float(turnover_match.group(1).replace(",", ""))
        actual = mean([float(credentials.gross_fee_receipts_fy1 or 0), float(credentials.gross_fee_receipts_fy2 or 0), float(credentials.gross_fee_receipts_fy3 or 0)])
        criteria.append({"criterion": "Average gross fee receipts", "required_value": required, "firm_value": round(actual, 2), "eligible": actual >= required, "gap": None if actual >= required else round(required - actual, 2)})
    staff_match = re.search(r"(?:minimum|at least)\s+(\d+)\s+(?:staff|professionals|team members)", text or "", re.I)
    if staff_match:
        required = int(staff_match.group(1))
        actual = int(credentials.total_staff or 0)
        criteria.append({"criterion": "Team capacity", "required_value": required, "firm_value": actual, "eligible": actual >= required, "gap": None if actual >= required else required - actual})
    if re.search(r"peer\s+review", text or "", re.I):
        ok = str(credentials.peer_review_status or "").lower() in ("valid", "active", "completed")
        criteria.append({"criterion": "Peer review status", "required_value": "Valid", "firm_value": credentials.peer_review_status or "Not provided", "eligible": ok, "gap": None if ok else "Valid peer review required"})
    if not criteria:
        criteria.append({"criterion": "Credential review", "required_value": "Firm profile available", "firm_value": credentials.firm_name, "eligible": True, "gap": None})
    gaps = [str(row["gap"]) for row in criteria if not row["eligible"]]
    return {"criteria": criteria, "overall_eligible": not gaps, "disqualifying_gaps": gaps, "recommended_approach": "Proceed with a credential-backed technical bid." if not gaps else "Resolve disqualifying gaps before submission."}


def generate_bid_proposal(title, rfp_text, eligibility, credentials):
    if not eligibility.get("overall_eligible"):
        return None
    partners = ", ".join(str(row.get("name")) for row in (credentials.partners or []) if row.get("name")) or "the engagement partner team"
    industries = ", ".join(str(row.get("industry")) for row in (credentials.industries_served or []) if row.get("industry")) or "relevant sectors"
    engagements = "\n".join(f"- {row.get('name', row.get('description', 'Relevant engagement'))}" for row in (credentials.key_engagements or [])) or "- Credential-backed engagements will be provided in the annexure."
    return f"""TECHNICAL BID: {title}

1. Executive Summary
{credentials.firm_name} is an ICAI-registered firm with registration number {credentials.icai_regn_no or 'to be inserted'}. The firm has {credentials.total_staff} staff members and experience across {industries}.

2. Understanding of the Assignment
We understand that the assignment requires a disciplined, evidence-based audit response aligned to the scope and eligibility conditions stated in the RFP.

3. Approach & Methodology
Our approach covers planning, risk assessment, fieldwork, control testing, substantive procedures, issue validation, partner review, and final reporting.

4. Team Composition
The proposed leadership team includes {partners}. Detailed role allocation will be finalized during mobilisation.

5. Relevant Experience
{engagements}

6. Quality Assurance Framework
The engagement will follow the firm's documented quality controls. Peer review status: {credentials.peer_review_status or 'to be confirmed'}.

Eligibility conclusion: all extracted criteria have been satisfied based solely on the stored firm credentials."""


def build_rows_xlsx(sheet_name, headers, rows):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = sheet_name[:31]
    fill = PatternFill("solid", fgColor="1F4E79")
    for col, header in enumerate(headers, 1):
        cell = ws.cell(1, col, header)
        cell.font = Font(color="FFFFFF", bold=True)
        cell.fill = fill
    for row_index, row in enumerate(rows, 2):
        for col_index, value in enumerate(row, 1):
            ws.cell(row_index, col_index, value)
    stream = io.BytesIO()
    wb.save(stream)
    return stream.getvalue()


def build_dp_pdf(client_name, bank_name, statement):
    stream = io.BytesIO()
    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(stream, pagesize=A4, rightMargin=0.5 * inch, leftMargin=0.5 * inch)
    story = [Paragraph("Drawing Power & Stock Statement", styles["Title"]), Paragraph(f"{client_name} | {bank_name} | {statement.period}", styles["Normal"]), Spacer(1, 12)]
    rows = [
        ["Particular", "Amount (INR)"],
        ["Gross stock", f"{float(statement.gross_stock):,.2f}"],
        ["Eligible stock", f"{float(statement.eligible_stock):,.2f}"],
        ["Gross debtors", f"{float(statement.gross_debtors):,.2f}"],
        ["Eligible debtors", f"{float(statement.eligible_debtors):,.2f}"],
        ["Creditors", f"{float(statement.creditors):,.2f}"],
        ["Drawing power", f"{float(statement.drawing_power):,.2f}"],
    ]
    table = Table(rows, colWidths=[3.8 * inch, 2.2 * inch])
    table.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F4E79")), ("TEXTCOLOR", (0, 0), (-1, 0), colors.white), ("GRID", (0, 0), (-1, -1), 0.5, colors.grey), ("ALIGN", (1, 1), (-1, -1), "RIGHT")]))
    story.append(table)
    doc.build(story)
    return stream.getvalue()


def profitability_summary(activities, entries):
    seconds = sum(row.duration_seconds or 0 for row in activities)
    billable_seconds = sum(row.duration_seconds or 0 for row in activities if ACTIVITY_BILLING_MAP.get(row.activity_type, {}).get("billable"))
    logged = sum(float(row.hours_logged or 0) for row in entries)
    revenue = sum(float(row.hours_logged or 0) * float(row.billing_rate or 0) for row in entries if row.billable)
    cost = sum(float(row.hours_logged or 0) * float(row.cost_rate or 0) for row in entries)
    breakdown = defaultdict(float)
    for row in activities:
        category = ACTIVITY_BILLING_MAP.get(row.activity_type, {"type": row.activity_type})["type"]
        breakdown[category] += round((row.duration_seconds or 0) / 3600, 2)
    actual = round(seconds / 3600, 2)
    return {
        "actual_hours": actual, "logged_hours": round(logged, 2),
        "billable_hours": round(billable_seconds / 3600, 2), "variance_hours": round(actual - logged, 2),
        "utilization_pct": round(billable_seconds / max(seconds, 1) * 100, 1),
        "revenue": round(revenue, 2), "cost": round(cost, 2), "margin": round(revenue - cost, 2),
        "task_breakdown": dict(breakdown),
    }
