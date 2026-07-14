"""AI Audit Working Papers engine."""
import io
import json
import logging
from datetime import datetime

import pandas as pd

logger = logging.getLogger(__name__)


def parse_trial_balance(df: pd.DataFrame) -> dict:
    rename_map = {
        "Account": "account", "Group": "group",
        "Debit": "debit", "Credit": "credit",
        "Dr": "debit", "Cr": "credit",
    }
    df = df.rename(columns=rename_map)
    if "account" not in df.columns:
        df["account"] = [f"Account {idx + 1}" for idx in range(len(df))]
    if "group" not in df.columns:
        df["group"] = ""
    df["debit"]  = pd.to_numeric(df.get("debit", 0), errors="coerce").fillna(0)
    df["credit"] = pd.to_numeric(df.get("credit", 0), errors="coerce").fillna(0)
    df["net"] = df["debit"] - df["credit"]

    def group_sum(grp_name):
        if "group" not in df.columns:
            return 0.0
        return float(df[df["group"].str.lower().str.contains(grp_name, na=False)]["net"].sum())

    return {
        "total_assets":       group_sum("asset"),
        "total_liabilities":  abs(group_sum("liabilit")),
        "revenue":            abs(group_sum("revenue")),
        "expenses":           group_sum("expense"),
        "current_assets":     group_sum("current asset"),
        "current_liabilities": abs(group_sum("current liabilit")),
        "accounts": df.to_dict("records"),
    }


def generate_audit_observations_fallback(tb: dict, ratios: dict, anomalies: list) -> str:
    """Create deterministic working-paper observations when LLM providers are unavailable."""
    observations = []
    if ratios.get("current_ratio", 0) < 1:
        observations.append((
            "Liquidity pressure",
            "balance sheet",
            "High",
            "Current liabilities exceed current assets based on the uploaded trial balance.",
            "There may be short-term liquidity stress or classification errors.",
            "Verify ageing schedules, bank facilities, subsequent payments, and current/non-current classification.",
        ))
    if ratios.get("debt_equity_ratio", 0) > 2:
        observations.append((
            "High leverage",
            "balance sheet",
            "Medium",
            "Debt-equity ratio is elevated for the period under review.",
            "Borrowing costs, covenant compliance, and going-concern assumptions may need closer review.",
            "Check loan confirmations, sanction letters, repayment schedules, and interest accruals.",
        ))
    if ratios.get("gross_margin_pct", 0) < 10:
        observations.append((
            "Low gross margin",
            "P&L",
            "Medium",
            "Gross margin appears low based on revenue and expense grouping.",
            "Revenue recognition, purchase cut-off, or expense classification may be misstated.",
            "Perform analytical review against prior period, GST turnover, and purchase register.",
        ))
    for anomaly in anomalies[:3]:
        observations.append((
            str(anomaly.get("type", "Flagged anomaly")).replace("_", " ").title(),
            "internal control",
            "High" if float(anomaly.get("risk_score") or 0) >= 70 else "Medium",
            str(anomaly.get("details") or "A system anomaly was linked to this client."),
            "The item may affect audit risk assessment and substantive testing scope.",
            "Review source documents, management explanations, and approval trail.",
        ))
    if not observations:
        observations.append((
            "Standard analytical review",
            "overall financial statements",
            "Low",
            "No high-risk analytical exception was detected from the uploaded trial balance.",
            "Audit conclusions still depend on source evidence and professional judgement.",
            "Complete routine lead schedules, external confirmations, and material balance testing.",
        ))

    lines = []
    for title, area, risk, finding, implication, recommendation in observations[:8]:
        lines.extend([
            f"OBSERVATION: {title}",
            f"AREA: {area}",
            f"RISK LEVEL: {risk}",
            f"FINDING: {finding}",
            f"IMPLICATION: {implication}",
            f"RECOMMENDATION: {recommendation}",
            "",
        ])
    return "\n".join(lines).strip()


def compute_ratios(tb: dict) -> dict:
    rev  = tb.get("revenue") or 1
    exp  = tb.get("expenses") or 0
    ast  = tb.get("total_assets") or 1
    liab = tb.get("total_liabilities") or 0
    cur_ast  = tb.get("current_assets") or 0
    cur_liab = tb.get("current_liabilities") or 1

    equity = ast - liab
    return {
        "gross_margin_pct":  round((rev - exp) / rev * 100, 2),
        "debt_equity_ratio": round(liab / max(equity, 0.01), 2),
        "current_ratio":     round(cur_ast / max(cur_liab, 0.01), 2),
        "asset_turnover":    round(rev / ast, 2),
    }


def generate_audit_observations(tb: dict, ratios: dict, anomalies: list,
                                 anthropic_client) -> str:
    prompt = f"""You are an audit assistant for an ICAI-registered Chartered Accountant.
Generate audit observations in standard ICAI working paper format.

Trial Balance Summary:
{json.dumps({k: v for k, v in tb.items() if k != 'accounts'}, default=str)}

Key Financial Ratios:
{json.dumps(ratios)}

Flagged Anomalies:
{json.dumps(anomalies[:10])}

For each observation use this EXACT structure:
OBSERVATION: [brief title]
AREA: [balance sheet / P&L / internal control]
RISK LEVEL: [High / Medium / Low]
FINDING: [what was found]
IMPLICATION: [potential impact]
RECOMMENDATION: [what the CA should verify or report]

Generate 5-8 observations. Focus on material items only."""

    message = anthropic_client.messages.create(
        model="claude-sonnet-4-20250514",
        max_tokens=3000,
        messages=[{"role": "user", "content": prompt}],
    )
    return message.content[0].text


def export_working_paper(client_name: str, observations: str,
                         ratios: dict, period: str) -> bytes:
    """Build a DOCX working paper and return its bytes."""
    from docx import Document as DocxDoc
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn

    doc = DocxDoc()
    doc.add_heading(f"Audit Working Paper — {client_name}", 0)
    doc.add_paragraph(f"Period: {period}")
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%d %b %Y')}")
    doc.add_paragraph("")

    doc.add_heading("Key Financial Ratios", level=1)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.rows[0].cells[0].text = "Ratio"
    tbl.rows[0].cells[1].text = "Value"
    for k, v in ratios.items():
        row = tbl.add_row()
        row.cells[0].text = k.replace("_", " ").title()
        row.cells[1].text = str(v)

    doc.add_paragraph("")
    doc.add_heading("Audit Observations", level=1)
    doc.add_paragraph(observations)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
